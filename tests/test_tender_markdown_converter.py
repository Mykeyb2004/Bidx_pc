from pathlib import Path
from types import SimpleNamespace

import openpyxl
import pytest
from docx import Document

from bid_writer.tender_markdown_converter import TenderConversionError, convert_tender_document


def test_converts_docx_headings_paragraphs_and_tables(tmp_path: Path):
    doc_path = tmp_path / "tender.docx"
    document = Document()
    document.add_heading("项目采购需求", level=1)
    document.add_paragraph("服务内容包括调查、分析、成果提交和验收。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "服务范围"
    table.cell(1, 1).text = "覆盖采购人指定区域"
    document.add_heading("评分标准", level=1)
    document.add_paragraph("评分总分为100分。")
    document.save(doc_path)

    result = convert_tender_document(doc_path, tmp_path / "out")

    markdown = result.converted_markdown_path.read_text(encoding="utf-8")
    assert "# 项目采购需求" in markdown
    assert "服务内容包括调查" in markdown
    assert "| 项目 | 要求 |" in markdown
    assert "# 评分标准" in markdown
    assert result.conversion_map_path.exists()
    assert any(block.block_type == "table" for block in result.blocks)


def test_converts_xlsx_sheets_to_markdown_tables(tmp_path: Path):
    workbook_path = tmp_path / "score.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "技术商务评分表"
    sheet["A1"] = "评分项"
    sheet["B1"] = "评分标准"
    sheet["C1"] = "分值"
    sheet["A2"] = "服务方案"
    sheet["B2"] = "完整得20分"
    sheet["C2"] = "20分"
    workbook.save(workbook_path)

    result = convert_tender_document(workbook_path, tmp_path / "out")

    markdown = result.converted_markdown_path.read_text(encoding="utf-8")
    assert "## 工作表：技术商务评分表" in markdown
    assert "| 评分项 | 评分标准 | 分值 |" in markdown
    assert any(block.sheet_name == "技术商务评分表" for block in result.blocks)


def test_rejects_wps_and_et_formats(tmp_path: Path):
    wps_path = tmp_path / "tender.wps"
    wps_path.write_text("fake", encoding="utf-8")

    with pytest.raises(TenderConversionError) as exc:
        convert_tender_document(wps_path, tmp_path / "out")

    assert "暂不支持 WPS 原生格式" in str(exc.value)


def test_pdf_without_text_layer_raises_clear_error(monkeypatch, tmp_path: Path):
    pdf_path = tmp_path / "scan.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 fake")

    class FakePage:
        def get_text(self, _kind):
            return ""

    class FakeDoc(list):
        def __init__(self):
            super().__init__([FakePage(), FakePage()])

    monkeypatch.setitem(__import__("sys").modules, "fitz", SimpleNamespace(open=lambda _path: FakeDoc()))
    monkeypatch.setitem(__import__("sys").modules, "pymupdf4llm", SimpleNamespace(to_markdown=lambda _path: ""))

    with pytest.raises(TenderConversionError) as exc:
        convert_tender_document(pdf_path, tmp_path / "out")

    assert "暂不支持 OCR" in str(exc.value)


def test_pdf_conversion_disables_ocr_and_records_page_numbers(monkeypatch, tmp_path: Path):
    pdf_path = tmp_path / "text.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 fake")
    captured = {}

    class FakePage:
        def __init__(self, text: str):
            self.text = text

        def get_text(self, _kind):
            return self.text

    class FakeDoc(list):
        def __init__(self):
            super().__init__(
                [
                    FakePage("项目采购需求 " + "服务内容" * 20),
                    FakePage("评分标准 " + "评审分值" * 20),
                ]
            )

    def fake_to_markdown(path, **kwargs):
        captured["path"] = path
        captured["kwargs"] = kwargs
        return [
            {"metadata": {"page": 1}, "text": "# 项目采购需求\n\n服务内容。"},
            {"metadata": {"page": 2}, "text": "# 评分标准\n\n| 评分项 | 分值 |\n| --- | --- |\n| 服务 | 10分 |"},
        ]

    monkeypatch.setitem(__import__("sys").modules, "fitz", SimpleNamespace(open=lambda _path: FakeDoc()))
    monkeypatch.setitem(__import__("sys").modules, "pymupdf4llm", SimpleNamespace(to_markdown=fake_to_markdown))

    result = convert_tender_document(pdf_path, tmp_path / "out")

    assert captured["path"] == str(pdf_path.resolve())
    assert captured["kwargs"]["use_ocr"] is False
    assert captured["kwargs"]["page_chunks"] is True
    assert {block.page_number for block in result.blocks} == {1, 2}
    assert any(block.heading_title == "项目采购需求" for block in result.blocks)


def test_doc_requires_libreoffice_when_no_converter_available(monkeypatch, tmp_path: Path):
    doc_path = tmp_path / "old.doc"
    doc_path.write_text("fake", encoding="utf-8")

    monkeypatch.setattr("bid_writer.tender_markdown_converter.shutil.which", lambda _name: None)

    with pytest.raises(TenderConversionError) as exc:
        convert_tender_document(doc_path, tmp_path / "out")

    assert "未检测到 LibreOffice" in str(exc.value)


def test_xls_uses_calamine_before_libreoffice(monkeypatch, tmp_path: Path):
    xls_path = tmp_path / "old.xls"
    xls_path.write_text("fake", encoding="utf-8")

    class FakeSheet:
        def to_python(self):
            return [["评分项", "评分标准", "分值"], ["服务方案", "完整得20分", "20分"]]

    class FakeWorkbook:
        sheet_names = ["技术商务评分表"]

        @classmethod
        def from_path(cls, path):
            assert path == str(xls_path.resolve())
            return cls()

        def get_sheet_by_name(self, sheet_name):
            assert sheet_name == "技术商务评分表"
            return FakeSheet()

    monkeypatch.setitem(
        __import__("sys").modules,
        "python_calamine",
        SimpleNamespace(CalamineWorkbook=FakeWorkbook),
    )
    monkeypatch.setattr("bid_writer.tender_markdown_converter.shutil.which", lambda _name: None)

    result = convert_tender_document(xls_path, tmp_path / "out")

    markdown = result.converted_markdown_path.read_text(encoding="utf-8")
    assert "## 工作表：技术商务评分表" in markdown
    assert "| 评分项 | 评分标准 | 分值 |" in markdown
    assert result.warnings == ("已通过 python-calamine 读取 .xls。",)
    assert any(block.cell_range == "A1:C2" for block in result.blocks)


def test_xlsx_expands_merged_cells_and_records_actual_cell_range(tmp_path: Path):
    workbook_path = tmp_path / "merged.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "评分表"
    sheet.merge_cells("B2:D2")
    sheet["B2"] = "评分标准"
    sheet["B3"] = "评分项"
    sheet["C3"] = "评审内容"
    sheet["D3"] = "分值"
    sheet["B4"] = "服务方案"
    sheet["C4"] = "完整得20分"
    sheet["D4"] = "20分"
    workbook.save(workbook_path)

    result = convert_tender_document(workbook_path, tmp_path / "out")

    markdown = result.converted_markdown_path.read_text(encoding="utf-8")
    table_block = next(block for block in result.blocks if block.block_type == "table")
    assert "评分标准" in markdown
    assert table_block.cell_range == "B2:D4"


def test_libreoffice_conversion_reports_subprocess_failure(monkeypatch, tmp_path: Path):
    doc_path = tmp_path / "old.doc"
    doc_path.write_text("fake", encoding="utf-8")

    monkeypatch.setattr("bid_writer.tender_markdown_converter.shutil.which", lambda _name: "/usr/bin/soffice")
    monkeypatch.setattr(
        "bid_writer.tender_markdown_converter.subprocess.run",
        lambda *_args, **_kwargs: SimpleNamespace(returncode=1, stderr="bad input", stdout=""),
    )

    with pytest.raises(TenderConversionError) as exc:
        convert_tender_document(doc_path, tmp_path / "out")

    assert "LibreOffice 转换失败：bad input" in str(exc.value)
