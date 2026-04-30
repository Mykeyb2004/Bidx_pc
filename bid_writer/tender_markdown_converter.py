"""招标文件到 Markdown block 的转换器。"""

from __future__ import annotations

import json
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path

import openpyxl
from docx import Document

from .tender_import_models import ConvertedBlock, TenderConversionResult, dump_conversion_map


class TenderConversionError(RuntimeError):
    """招标文件转换失败。"""


UNSUPPORTED_WPS_SUFFIXES = {".wps", ".et"}
SUPPORTED_SUFFIXES = {".pdf", ".docx", ".xlsx", ".doc", ".xls"} | UNSUPPORTED_WPS_SUFFIXES


@dataclass(frozen=True)
class SheetTable:
    rows: list[list[str]]
    cell_range: str


def convert_tender_document(path: Path, output_dir: Path) -> TenderConversionResult:
    source_path = Path(path).expanduser().resolve()
    suffix = source_path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise TenderConversionError(f"暂不支持该文件格式：{suffix or '无扩展名'}")
    if suffix in UNSUPPORTED_WPS_SUFFIXES:
        raise TenderConversionError("暂不支持 WPS 原生格式，请另存为 .docx、.xlsx 或可复制文字 PDF 后再导入。")

    output_dir = Path(output_dir).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    source_for_conversion = source_path
    warnings: list[str] = []

    if suffix == ".doc":
        source_for_conversion = _convert_with_libreoffice(source_path, output_dir, "docx")
        suffix = ".docx"
        warnings.append("已通过 LibreOffice 将 .doc 预转换为 .docx。")
    elif suffix == ".xls":
        calamine_blocks = _try_convert_xls_with_calamine(source_path, source_path.name)
        if calamine_blocks is not None:
            return _write_conversion_result(
                source_path=source_path,
                output_dir=output_dir,
                blocks=calamine_blocks,
                warnings=("已通过 python-calamine 读取 .xls。",),
            )
        source_for_conversion = _convert_with_libreoffice(source_path, output_dir, "xlsx")
        suffix = ".xlsx"
        warnings.append("已通过 LibreOffice 将 .xls 预转换为 .xlsx。")

    if suffix == ".docx":
        blocks = _convert_docx(source_for_conversion, source_path.name)
    elif suffix == ".xlsx":
        blocks = _convert_xlsx(source_for_conversion, source_path.name)
    elif suffix == ".pdf":
        blocks = _convert_pdf(source_for_conversion, source_path.name)
    else:
        raise TenderConversionError(f"暂不支持该文件格式：{suffix}")

    if not blocks:
        raise TenderConversionError("未从文件中解析到可用文本。")

    return _write_conversion_result(
        source_path=source_path,
        output_dir=output_dir,
        blocks=blocks,
        warnings=tuple(warnings),
    )


def _write_conversion_result(
    *,
    source_path: Path,
    output_dir: Path,
    blocks: list[ConvertedBlock],
    warnings: tuple[str, ...],
) -> TenderConversionResult:
    converted_path = output_dir / "converted.md"
    map_path = output_dir / "conversion_map.json"
    converted_path.write_text(_join_blocks(blocks), encoding="utf-8")
    result = TenderConversionResult(
        source_path=source_path,
        output_dir=output_dir,
        converted_markdown_path=converted_path,
        conversion_map_path=map_path,
        blocks=blocks,
        warnings=warnings,
    )
    map_path.write_text(
        json.dumps(dump_conversion_map(result), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return result


def _convert_docx(path: Path, source_name: str) -> list[ConvertedBlock]:
    document = Document(path)
    blocks: list[ConvertedBlock] = []
    order = 0
    paragraph_index = 0
    table_index = 0

    for child in document.element.body:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = document.paragraphs[paragraph_index]
            paragraph_index += 1
            text = paragraph.text.strip()
            if not text:
                continue
            order += 1
            level = _docx_heading_level(paragraph.style.name if paragraph.style else "")
            markdown = f"{'#' * level} {text}" if level else text
            blocks.append(
                ConvertedBlock(
                    block_id=f"docx:p{paragraph_index:04d}",
                    source_file=source_name,
                    source_type="docx",
                    block_type="heading" if level else "paragraph",
                    markdown=markdown,
                    text=text,
                    order_index=order,
                    heading_level=level,
                    heading_title=text if level else "",
                    paragraph_index=paragraph_index,
                )
            )
        elif tag == "tbl":
            table = document.tables[table_index]
            table_index += 1
            markdown = _markdown_table([[cell.text.strip() for cell in row.cells] for row in table.rows])
            if not markdown.strip():
                continue
            order += 1
            blocks.append(
                ConvertedBlock(
                    block_id=f"docx:t{table_index:04d}",
                    source_file=source_name,
                    source_type="docx",
                    block_type="table",
                    markdown=markdown,
                    text=markdown,
                    order_index=order,
                    table_index=table_index,
                )
            )
    return blocks


def _docx_heading_level(style_name: str) -> int | None:
    match = re.search(r"(?:Heading|标题)\s*([1-6一二三四五六])", style_name, re.I)
    if not match:
        return None
    raw = match.group(1)
    chinese = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6}
    return chinese.get(raw, int(raw) if raw.isdigit() else 1)


def _convert_xlsx(path: Path, source_name: str) -> list[ConvertedBlock]:
    workbook = openpyxl.load_workbook(path, data_only=True)
    blocks: list[ConvertedBlock] = []
    order = 0
    for sheet in workbook.worksheets:
        order += 1
        sheet_heading = f"工作表：{sheet.title}"
        blocks.append(
            ConvertedBlock(
                block_id=f"xlsx:{sheet.title}:heading",
                source_file=source_name,
                source_type="xlsx",
                block_type="heading",
                markdown=f"## {sheet_heading}",
                text=sheet_heading,
                order_index=order,
                heading_level=2,
                heading_title=sheet_heading,
                sheet_name=sheet.title,
            )
        )
        table = _sheet_table(sheet)
        if table.rows:
            order += 1
            markdown = _markdown_table(table.rows)
            blocks.append(
                ConvertedBlock(
                    block_id=f"xlsx:{sheet.title}:table1",
                    source_file=source_name,
                    source_type="xlsx",
                    block_type="table",
                    markdown=markdown,
                    text=markdown,
                    order_index=order,
                    sheet_name=sheet.title,
                    cell_range=table.cell_range,
                    table_index=1,
                )
            )
    return blocks


def _try_convert_xls_with_calamine(path: Path, source_name: str) -> list[ConvertedBlock] | None:
    try:
        from python_calamine import CalamineWorkbook
    except ImportError:
        return None
    try:
        workbook = CalamineWorkbook.from_path(str(path))
    except Exception:
        return None
    blocks: list[ConvertedBlock] = []
    order = 0
    for sheet_name in workbook.sheet_names:
        sheet = workbook.get_sheet_by_name(sheet_name)
        rows = [[str(cell).strip() if cell is not None else "" for cell in row] for row in sheet.to_python()]
        rows = [row for row in rows if any(row)]
        order += 1
        heading = f"工作表：{sheet_name}"
        blocks.append(
            ConvertedBlock(
                block_id=f"xls:{sheet_name}:heading",
                source_file=source_name,
                source_type="xls",
                block_type="heading",
                markdown=f"## {heading}",
                text=heading,
                order_index=order,
                heading_level=2,
                heading_title=heading,
                sheet_name=sheet_name,
            )
        )
        if rows:
            rows = _normalize_rows(rows)
            order += 1
            markdown = _markdown_table(rows)
            blocks.append(
                ConvertedBlock(
                    block_id=f"xls:{sheet_name}:table1",
                    source_file=source_name,
                    source_type="xls",
                    block_type="table",
                    markdown=markdown,
                    text=markdown,
                    order_index=order,
                    sheet_name=sheet_name,
                    cell_range=_a1_range(1, 1, len(rows), max(len(row) for row in rows)),
                    table_index=1,
                )
            )
    return blocks


def _sheet_table(sheet) -> SheetTable:
    merged_lookup = _merged_cell_value_lookup(sheet)
    rows: list[tuple[int, list[tuple[int, str]]]] = []
    min_row: int | None = None
    max_row: int | None = None
    min_col: int | None = None
    max_col: int | None = None
    for row in sheet.iter_rows():
        values: list[tuple[int, str]] = []
        for cell in row:
            value = merged_lookup.get((cell.row, cell.column), cell.value)
            text = "" if value is None else str(value).strip()
            values.append((cell.column, text))
            if text:
                min_row = cell.row if min_row is None else min(min_row, cell.row)
                max_row = cell.row if max_row is None else max(max_row, cell.row)
                min_col = cell.column if min_col is None else min(min_col, cell.column)
                max_col = cell.column if max_col is None else max(max_col, cell.column)
        if any(text for _column, text in values):
            rows.append((row[0].row, values))
    if min_row is None or max_row is None or min_col is None or max_col is None:
        return SheetTable(rows=[], cell_range="")

    table_rows: list[list[str]] = []
    by_row = {row_number: dict(values) for row_number, values in rows}
    for row_number in range(min_row, max_row + 1):
        row_values = by_row.get(row_number, {})
        table_rows.append([row_values.get(column, "") for column in range(min_col, max_col + 1)])
    return SheetTable(
        rows=table_rows,
        cell_range=_a1_range(min_row, min_col, max_row, max_col),
    )


def _merged_cell_value_lookup(sheet) -> dict[tuple[int, int], object]:
    lookup: dict[tuple[int, int], object] = {}
    for cell_range in sheet.merged_cells.ranges:
        value = sheet.cell(cell_range.min_row, cell_range.min_col).value
        for row in range(cell_range.min_row, cell_range.max_row + 1):
            for column in range(cell_range.min_col, cell_range.max_col + 1):
                lookup[(row, column)] = value
    return lookup


def _normalize_rows(rows: list[list[str]]) -> list[list[str]]:
    width = max(len(row) for row in rows)
    return [row + [""] * (width - len(row)) for row in rows]




def _convert_pdf(path: Path, source_name: str) -> list[ConvertedBlock]:
    try:
        import fitz
        import pymupdf4llm
    except ImportError as exc:
        raise TenderConversionError(f"PDF 转换依赖未安装：{exc}") from exc

    doc = fitz.open(path)
    extracted_text = "\n".join(page.get_text("text") for page in doc)
    if _visible_text_length(extracted_text) < max(80, len(doc) * 20):
        raise TenderConversionError("该 PDF 可能没有可复制文本层；当前版本暂不支持 OCR。")
    page_chunks = pymupdf4llm.to_markdown(str(path), page_chunks=True, use_ocr=False)
    blocks: list[ConvertedBlock] = []
    order = 0
    for page_index, page_chunk in enumerate(_iter_pdf_page_chunks(page_chunks), start=1):
        page_number = _page_number(page_chunk, page_index)
        markdown = str(page_chunk.get("text", "")).strip()
        for chunk in _split_markdown_blocks(markdown):
            order += 1
            heading_level, heading_title = _markdown_heading(chunk)
            blocks.append(
                ConvertedBlock(
                    block_id=f"pdf:p{page_number:04d}:b{order:04d}",
                    source_file=source_name,
                    source_type="pdf",
                    block_type="heading" if heading_level else ("table" if "|" in chunk else "paragraph"),
                    markdown=chunk,
                    text=re.sub(r"^#+\s*", "", chunk).strip(),
                    order_index=order,
                    heading_level=heading_level,
                    heading_title=heading_title,
                    page_number=page_number,
                )
            )
    return blocks


def _iter_pdf_page_chunks(page_chunks) -> list[dict]:
    if isinstance(page_chunks, list):
        return [chunk for chunk in page_chunks if isinstance(chunk, dict)]
    return [{"metadata": {"page": 1}, "text": str(page_chunks)}]


def _page_number(page_chunk: dict, fallback: int) -> int:
    metadata = page_chunk.get("metadata") or {}
    raw = metadata.get("page") or metadata.get("page_number") or fallback
    try:
        page_number = int(raw)
    except (TypeError, ValueError):
        page_number = fallback
    return max(1, page_number)


def _split_markdown_blocks(markdown: str) -> list[str]:
    return [part.strip() for part in re.split(r"\n\s*\n", markdown) if part.strip()]


def _markdown_heading(markdown: str) -> tuple[int | None, str]:
    first_line = markdown.splitlines()[0].strip()
    match = re.match(r"^(#{1,6})\s+(.+)$", first_line)
    if not match:
        return None, ""
    return len(match.group(1)), match.group(2).strip()


def _visible_text_length(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", text))


def _markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * width
    body = normalized[1:]
    lines = [
        "| " + " | ".join(_escape_cell(cell) for cell in header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(_escape_cell(cell) for cell in row) + " |")
    return "\n".join(lines)


def _escape_cell(value: str) -> str:
    return str(value).replace("\n", "<br>").replace("|", "\\|")


def _a1_range(min_row: int, min_col: int, max_row: int, max_col: int) -> str:
    return f"{_column_letter(min_col)}{min_row}:{_column_letter(max_col)}{max_row}"


def _column_letter(index: int) -> str:
    letters = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        letters = chr(65 + remainder) + letters
    return letters or "A"


def _join_blocks(blocks: list[ConvertedBlock]) -> str:
    return "\n\n".join(block.markdown.strip() for block in blocks if block.markdown.strip()).strip() + "\n"


def _convert_with_libreoffice(path: Path, output_dir: Path, target_ext: str) -> Path:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if executable is None:
        raise TenderConversionError("未检测到 LibreOffice，无法转换旧 Office 格式；请另存为 .docx 或 .xlsx 后再导入。")
    command = [
        executable,
        "--headless",
        "--convert-to",
        target_ext,
        "--outdir",
        str(output_dir),
        str(path),
    ]
    completed = subprocess.run(command, check=False, capture_output=True, text=True)
    if completed.returncode != 0:
        message = completed.stderr.strip() or completed.stdout.strip()
        raise TenderConversionError(f"LibreOffice 转换失败：{message}")
    converted = output_dir / f"{path.stem}.{target_ext}"
    if not converted.exists():
        raise TenderConversionError("LibreOffice 转换完成但未找到输出文件。")
    return converted
